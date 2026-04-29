#!/usr/bin/env python3
"""
Medical PDF OCR + Romanian Translation Pipeline
 
"""

import argparse
import base64
import json
import sys
import time
from concurrent.futures import ThreadPoolExecutor, as_completed
from io import BytesIO
from pathlib import Path
from threading import Lock

import requests
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image, ImageStat, ImageFilter

# ─── Configuration ────────────────────────────────────────────────────────────

OLLAMA_URL      = "http://localhost:21434/api"
OCR_MODEL       = "qwen3-vl:8b"
TRANSLATE_MODEL = "translategemma:12b"      # Best for medical translation

DPI             = 300
TIMEOUT         = 60
MAX_IMG_SIZE    = 1536

OCR_WORKERS       = 1      # Keep at 1 for vision model
TRANSLATE_WORKERS = 2      # Safe to increase on strong GPU

# ─── Progress ────────────────────────────────────────────────────────────────

progress_lock = Lock()

def save_progress(progress_file: Path, data: dict):
    try:
        with progress_lock:
            with open(progress_file, "w", encoding="utf-8") as f:
                json.dump(data, f, ensure_ascii=False, indent=2)
        print(f"💾 Progress saved ({len(data)} pages)")
    except Exception as e:
        print(f"❌ Failed to save progress: {e}")


def load_progress(progress_file: Path) -> dict:
    if progress_file.exists():
        try:
            with open(progress_file, "r", encoding="utf-8") as f:
                return json.load(f)
        except:
            pass
    return {}


def is_blank_page(image: Image.Image, threshold: float = 0.85) -> bool:
    """Improved blank page detection"""
    gray = image.convert("L")
    stat = ImageStat.Stat(gray)
    mean = stat.mean[0] / 255.0
    variance = stat.var[0]
    
    edges = gray.filter(ImageFilter.FIND_EDGES)
    edge_mean = ImageStat.Stat(edges).mean[0] / 255.0

    return (mean > threshold and variance < 1200 and edge_mean < 0.10)


# ─── Ollama Helpers ───────────────────────────────────────────────────────────

def check_ollama():
    try:
        r = requests.get(f"{OLLAMA_URL}/tags", timeout=10)
        r.raise_for_status()
        print("✅ Ollama is running\n")
    except:
        print("❌ Ollama not running. Start with: ollama serve")
        sys.exit(1)


def image_to_base64(pil_image: Image.Image) -> str:
    w, h = pil_image.size
    if max(w, h) > MAX_IMG_SIZE:
        ratio = MAX_IMG_SIZE / max(w, h)
        pil_image = pil_image.resize((int(w * ratio), int(h * ratio)), Image.LANCZOS)
    buf = BytesIO()
    pil_image.save(buf, format="JPEG", quality=95)
    return base64.b64encode(buf.getvalue()).decode("utf-8")


def ocr_page(image_b64: str, page_num: int) -> str:
    prompt = "You are an expert medical OCR engine. Extract ALL visible text accurately. Output ONLY the text."
    payload = {
        "model": OCR_MODEL,
        "prompt": prompt,
        "images": [image_b64],
        "stream": False,
        "options": {"temperature": 0.0, "num_predict": 8192},
    }
    try:
        r = requests.post(f"{OLLAMA_URL}/generate", json=payload, timeout=TIMEOUT)
        r.raise_for_status()
        return r.json().get("response", "").strip() or "[empty page]"
    except Exception as e:
        return f"[OCR error page {page_num}]"


# ─── Translation ─────────────────────────────────────────────────────────────

def split_into_chunks(text: str, max_chars: int = 1800, overlap: int = 250) -> list:
    if not text.strip():
        return [text]
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    chunks = []
    current = ""
    for para in paragraphs:
        if len(current) + len(para) + 2 > max_chars and current:
            chunks.append(current.strip())
            overlap_text = current[-overlap:] if len(current) > overlap else current
            current = overlap_text + "\n\n" + para
        else:
            current += "\n\n" + para if current else para
    if current:
        chunks.append(current.strip())
    return chunks


def translate_text(text: str, page_num: int) -> str:
    if not text.strip() or text.startswith("["):
        return text

    chunks = split_into_chunks(text)
    translated_chunks = []
    print(f"   → Page {page_num}: Translating {len(chunks)} chunk(s)")

    for i, chunk in enumerate(chunks, 1):
        prompt = (
            "You are a professional medical translator (Spanish → Romanian).\n\n"
            "STRICT RULES:\n"
            "- Output ONLY the Romanian translation.\n"
            "- NEVER invent or add information.\n"
            "- For abbreviations: Keep original + (full Romanian term)\n"
            "  Example: ITU (infecție de tract urinar), FRD (fosa renal dreaptă)\n"
            "- Preserve structure and line breaks.\n\n"
            f"Text:\n{chunk}"
        )
        payload = {
            "model": TRANSLATE_MODEL,
            "prompt": prompt,
            "stream": False,
            "options": {"temperature": 0.0},
        }
        try:
            r = requests.post(f"{OLLAMA_URL}/generate", json=payload, timeout=TIMEOUT)
            r.raise_for_status()
            translated_chunks.append(r.json().get("response", "").strip())
        except:
            translated_chunks.append(f"[Translation error chunk {i}]")

    return "\n\n".join(translated_chunks)


def process_ocr_job(job: dict):
    page_num = job["page_num"]
    image = job["image"]
    t0 = time.time()
    b64 = image_to_base64(image)
    ocr_text = ocr_page(b64, page_num)
    return {"page_num": page_num, "ocr": ocr_text, "time": time.time() - t0}


def build_docx(pages_data: dict, output_path: Path, pdf_name: str):
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)

    doc.add_heading(f"Traducere: {pdf_name}", level=1)
    doc.add_paragraph(f"OCR: {OCR_MODEL} | Translation: {TRANSLATE_MODEL} | DPI: {DPI}")
    doc.add_page_break()

    for p in sorted(pages_data.keys(), key=int):
        page_num = int(p)
        doc.add_heading(f"Pagina {page_num}", level=2)
        text = pages_data[p].get("translation", "[lipsă]")
        for line in text.split("\n"):
            if line.strip():
                doc.add_paragraph(line.strip())
        doc.add_page_break()

    doc.save(str(output_path))
    print(f"✅ Document saved: {output_path.name}")


# ─── Main ─────────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(description="Medical PDF OCR + Translation")
    parser.add_argument("--input", required=True)
    parser.add_argument("--start-page", type=int, default=1)
    parser.add_argument("--ocr-only", action="store_true")
    parser.add_argument("--skip-blank-pages", action="store_true", default=True)
    parser.add_argument("--blank-threshold", type=float, default=0.99)
    args = parser.parse_args()

    pdf_path = Path(args.input).resolve()
    stem = pdf_path.stem
    progress_file = pdf_path.parent / f"{stem}_progress.json"
    output_docx = pdf_path.parent / f"{stem}_romanian.docx"

    print(f"\n{'─'*90}")
    print("  MEDICAL PDF OCR + ROMANIAN TRANSLATION (Vast.ai Optimized)")
    print(f"{'─'*90}")
    print(f"  OCR Model       : {OCR_MODEL}")
    print(f"  Translation     : {TRANSLATE_MODEL}")
    print(f"  Skip blank      : {args.skip_blank_pages}")
    print(f"  Input           : {pdf_path.name}")
    print(f"{'─'*90}\n")

    check_ollama()
    pages_data = load_progress(progress_file)

    from pdf2image import convert_from_path
    images = convert_from_path(str(pdf_path), dpi=DPI, first_page=args.start_page, fmt="jpeg", thread_count=4)
    total_pages = args.start_page + len(images) - 1
    print(f"✅ Loaded {len(images)} pages\n")

    # OCR Phase
    print("🔍 Starting OCR Phase...\n")
    ocr_start = time.time()

    for i, img in enumerate(images):
        page_num = args.start_page + i
        key = str(page_num)

        if key in pages_data:
            continue

        if args.skip_blank_pages and is_blank_page(img, args.blank_threshold):
            print(f"⏭️  Page {page_num:>3}/{total_pages}  → Skipped (blank)")
            pages_data[key] = {"page": page_num, "ocr": "[blank page]", "translation": "[blank page]"}
            save_progress(progress_file, pages_data)
            continue

        res = process_ocr_job({"page_num": page_num, "image": img})
        print(f"✅ Page {page_num:>3}/{total_pages}  OCR {res['time']:5.1f}s")

        pages_data[key] = {"page": page_num, "ocr": res["ocr"], "translation": res["ocr"]}
        save_progress(progress_file, pages_data)

    print(f"\n⏱️  OCR Phase completed in {(time.time() - ocr_start)/60:.1f} minutes\n")

    # Translation Phase
    if not args.ocr_only:
        print("🌍 Starting Translation Phase...\n")
        tr_start = time.time()
        for key in sorted(pages_data.keys(), key=int):
            if pages_data[key].get("translation") == pages_data[key].get("ocr"):
                translation = translate_text(pages_data[key]["ocr"], int(key))
                pages_data[key]["translation"] = translation
                save_progress(progress_file, pages_data)
        print(f"⏱️  Translation completed in {(time.time()-tr_start)/60:.1f} minutes\n")

    build_docx(pages_data, output_docx, stem)
    print(f"\n🎉 ALL DONE! Output: {output_docx.name}")


if __name__ == "__main__":
    main()
